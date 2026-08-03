"""Document generation, filename construction, and spreadsheet helpers.

Provides the functions that produce the final ``.docx`` letter files,
the Excel control spreadsheet, and the safe Windows filenames for each
generated document.

Public exports:
    formatar_lista_pt: Format a Python list as a Portuguese enumeration string.
    frases_propositura: Return plural-aware phrase fragments for the letter template.
    normalizar_numero_mocao: Strip year suffixes from a motion number string.
    construir_nome_arquivo: Build a safe Windows filename for one letter.
    criar_modelo_planilha: Create (or overwrite) the Excel template with headers.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Any

__all__ = [
    "formatar_lista_pt",
    "frases_propositura",
    "normalizar_numero_mocao",
    "construir_nome_arquivo",
    "criar_modelo_planilha",
    "criar_modelo_envelope",
    "gerar_envelope_combinado",
    "ajustar_posicao_rodape",
    "remover_quebras_manuais",
]

# ── Pre-compiled regex patterns ───────────────────────────────────────────────

# Matches a year suffix that the AI may append to a motion number.
# Examples: "124/2026" → strips "/2026"; "124-26" → strips "-26".
_RE_ANO_MOCAO: re.Pattern[str] = re.compile(r"[-/]\d{2,4}$")

# Characters that are illegal in Windows file and folder names.
_RE_NOME_INVALIDO: re.Pattern[str] = re.compile(r'[\\/*?:"<>|]')

# Portuguese vowels (all lowercase) — used to detect likely acronyms.
_VOGAIS: frozenset[str] = frozenset("aeiouáéíóúâêôãõà")

# Portuguese prepositions and articles that should be lowercased mid-name.
_PREPS_PT: frozenset[str] = frozenset({
    "a", "ao", "aos", "à", "às",
    "com", "da", "das", "de", "do", "dos",
    "e", "em",
    "na", "nas", "no", "nos",
    "o", "os",
    "para", "por",
})

# Common Portuguese acronyms and abbreviations containing vowels to be kept uppercase.
_ACRONIMOS_PT: frozenset[str] = frozenset({
    "oab", "apae", "mds", "sus", "ong", "ongs", "eua", "onu", "cras", "creas",
    "ubs", "upa", "mei", "eireli", "ltda", "ead", "mec", "pt", "psdb", "pdt",
    "pl", "psol", "mdb", "psd", "pp", "pode", "pcd"
})


def _titlecase_nome(nome: str) -> str:
    """Convert a recipient name to title case for use in filenames.

    Tokens are preserved as all-caps when they:

    - Contain a period (abbreviations like ``S.A.`` or ``A.P.A.E.``)
    - Have no vowels (consonant-cluster abbreviations like ``BNB``, ``CNPJ``)
    - Match a predefined set of common Portuguese acronyms (``APAE``, ``OAB``, ``SUS``)

    Known Portuguese prepositions/articles are lowercased (except at
    position 0).  All other tokens are capitalised.
    """
    words = nome.split()
    if not words:
        return nome
    result: list[str] = []
    for i, word in enumerate(words):
        w_lower = word.lower()
        if "." in word:
            # Contains a period → abbreviation (S.A., A.P.A.E., Jr.)
            result.append(word)
        elif not any(c in _VOGAIS for c in w_lower):
            # No vowels → consonant-only abbreviation (BNB, CNPJ, SP)
            result.append(word)
        elif w_lower in _ACRONIMOS_PT:
            # Common acronyms containing vowels -> keep uppercase
            result.append(word.upper())
        elif i > 0 and w_lower in _PREPS_PT:
            # Mid-name preposition/article → lowercase
            result.append(w_lower)
        else:
            result.append(word.capitalize())
    return " ".join(result)


def formatar_lista_pt(items: list[str]) -> str:
    """Format a list of strings as a Portuguese enumeration, deduplicating order.

    Examples:
        ``["a"]`` → ``"a"``
        ``["a", "b"]`` → ``"a e b"``
        ``["a", "b", "c"]`` → ``"a, b e c"``

    Args:
        items: Strings to format.  Duplicates are removed while preserving order.

    Returns:
        Single formatted string.
    """
    unique: list[str] = list(dict.fromkeys(items))
    if len(unique) == 1:
        return unique[0]
    return ", ".join(unique[:-1]) + " e " + unique[-1]


def frases_propositura(
    tipo_propositura: str,
    tipo_mocao_merged: str,
    n_props: int,
) -> tuple[str, str, str]:
    """Return plural-aware phrase fragments for the letter template.

    Args:
        tipo_propositura: ``"mocao"`` or ``"requerimento_pesar"``.
        tipo_mocao_merged: Merged motion type string (e.g. ``"Aplauso"``).
            Ignored when *tipo_propositura* is ``"requerimento_pesar"``.
        n_props: Number of propositions grouped in this letter.

    Returns:
        A three-tuple ``(designacao_propositura, copia_art, aprovada_s)``
        where:

        - *designacao_propositura* — full noun phrase, e.g. ``"Moção de Aplauso"``
          or ``"Moções de Aplauso"``.
        - *copia_art* — contracted article phrase, e.g. ``"cópia da"`` or
          ``"cópias das"``.
        - *aprovada_s* — past-participle agreement, ``"aprovada"`` /
          ``"aprovadas"`` / ``"aprovado"`` / ``"aprovados"``.
    """
    if tipo_propositura == "requerimento_pesar":
        if n_props > 1:
            return "Requerimentos de Pesar", "cópias dos", "aprovados"
        return "Requerimento de Pesar", "cópia do", "aprovado"
    if n_props > 1:
        return f"Moções de {tipo_mocao_merged}", "cópias das", "aprovadas"
    return f"Moção de {tipo_mocao_merged}", "cópia da", "aprovada"


def normalizar_numero_mocao(numero: str) -> str:
    """Strip year suffixes that the AI may include in the motion number.

    Examples:
        ``"124/2026"`` → ``"124"``
        ``"124-26"`` → ``"124"``
        ``"124"`` → ``"124"`` (unchanged)

    Args:
        numero: Raw motion number string from the AI response.

    Returns:
        Normalised motion number string.
    """
    return _RE_ANO_MOCAO.sub("", numero).strip()


def construir_nome_arquivo(
    num_oficio_str: str,
    sigla_servidor: str,
    tipo_mocao: str,
    num_mocao: str,
    envio: str,
    nome_dest: str,
    sigla_autores: str,
    ano: int,
    tipo_propositura: str = "mocao",
) -> str:
    """Build a safe Windows filename for a generated letter document.

    For moções the filename format is::

        Of. {num} - {sigla} - Moção de {tipo} nº {num_mocao}-{yy} - {envio} - {dest} - {autores}.docx

    For requerimentos de pesar::

        Of. {num} - {sigla} - Req. de Pesar nº {num_mocao}-{yy} - {envio} - {dest} - {autores}.docx

    All characters that are invalid in Windows filenames are removed.

    Args:
        num_oficio_str: Zero-padded letter number (e.g. ``"001"``).
        sigla_servidor: Drafter's initials (e.g. ``"ajc"``).
        tipo_mocao: Motion type (e.g. ``"Aplauso"``).  Unused when
            *tipo_propositura* is ``"requerimento_pesar"``.
        num_mocao: Normalised propositura number (e.g. ``"124"``).
        envio: Delivery method (e.g. ``"E-mail"``).
        nome_dest: Recipient name as it appears in the address block.
        sigla_autores: Author sigla or combined sigla (e.g. ``"ad e outros"``).
        ano: Four-digit year of the propositura.
        tipo_propositura: Either ``"mocao"`` (default) or
            ``"requerimento_pesar"``.

    Returns:
        Sanitised filename string ending in ``.docx``.
    """
    nome_dest = _titlecase_nome(nome_dest)
    # Trim long recipient names so the full path stays under 240 chars on Windows.
    _MAX_DEST = 60
    nome_dest_trim = nome_dest[:_MAX_DEST].rstrip() if len(nome_dest) > _MAX_DEST else nome_dest

    ano_2d = f"{ano % 100:02d}"
    if tipo_propositura == "requerimento_pesar":
        nome = (
            f"Of. {num_oficio_str} - {sigla_servidor} - "
            f"Req. Pesar nº {num_mocao}-{ano_2d} - "
            f"{envio.lower()} - {nome_dest_trim} - {sigla_autores}.docx"
        )
    else:
        nome = (
            f"Of. {num_oficio_str} - {sigla_servidor} - "
            f"Moção de {tipo_mocao} nº {num_mocao}-{ano_2d} - "
            f"{envio.lower()} - {nome_dest_trim} - {sigla_autores}.docx"
        )
    return _RE_NOME_INVALIDO.sub("", nome)


def criar_modelo_planilha(destino: "str | Path | None" = None) -> Path:
    """Create the Excel spreadsheet template with formatted headers.

    If *destino* is ``None``, the file is placed alongside the executable
    (frozen mode) or the project root (dev mode).

    Args:
        destino: Target file path.  ``None`` uses the default location.

    Returns:
        Path of the file created.

    Raises:
        ImportError: If ``openpyxl`` is not installed.
    """
    from openpyxl import Workbook  # noqa: PLC0415 — lazy to avoid test startup cost
    from openpyxl.styles import Alignment, Font, PatternFill  # noqa: PLC0415

    if destino is None:
        from z7_officeletters.constants import MODELO_PLANILHA  # noqa: PLC0415

        if getattr(sys, "frozen", False):
            destino = Path(sys.executable).parent / MODELO_PLANILHA
        else:
            destino = Path(__file__).parent.parent.parent.parent / MODELO_PLANILHA
    destino = Path(destino)
    destino.parent.mkdir(parents=True, exist_ok=True)

    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "Controle"

    cabecalhos = ["Of. n.º", "Data", "Destinatário", "Assunto", "Vereador", "Envio", "Autor"]
    ws.append(cabecalhos)

    fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    fonte = Font(bold=True, color="FFFFFF", size=11)
    alin = Alignment(horizontal="center", vertical="center", wrap_text=False)

    for cell in ws[1]:
        cell.fill = fill
        cell.font = fonte
        cell.alignment = alin

    larguras: dict[str, int] = {
        "A": 10, "B": 12, "C": 32, "D": 54, "E": 32, "F": 14, "G": 10
    }
    for col, width in larguras.items():
        ws.column_dimensions[col].width = width

    ws.row_dimensions[1].height = 22
    ws.freeze_panes = "A2"

    wb.save(str(destino))
    return destino


def criar_modelo_envelope(destino: str | Path | None = None) -> Path:
    """Create the default Word envelope template with placeholders and DL size.

    If *destino* is ``None``, the file is placed in the default templates folder.

    Args:
        destino: Target file path.  ``None`` uses the default location.

    Returns:
        Path of the file created.
    """
    from docx import Document  # type: ignore[import-untyped]  # noqa: PLC0415
    from docx.shared import Cm, Pt  # type: ignore[import-untyped]  # noqa: PLC0415

    if destino is None:
        from z7_officeletters.constants import MODELO_ENVELOPE  # noqa: PLC0415

        if getattr(sys, "frozen", False):
            destino = Path(sys.executable).parent / MODELO_ENVELOPE
        else:
            destino = Path(__file__).parent.parent.parent.parent / MODELO_ENVELOPE
    destino = Path(destino)
    destino.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Set page size to DL envelope (22 cm x 11 cm in landscape orientation)
    section = doc.sections[0]
    section.page_width = Cm(22)
    section.page_height = Cm(11)

    # Set margins
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Add Recipient name aligned to the left (no treatment/honorific prefix)
    p_dest1 = doc.add_paragraph()
    p_dest1.paragraph_format.space_before = Pt(36)  # Add some spacing from the top
    p_dest1.paragraph_format.space_after = Pt(0)
    p_dest1.paragraph_format.line_spacing = 1.15

    r_dest1 = p_dest1.add_run("{{ DESTINATARIO_NOME }}")
    r_dest1.font.size = Pt(11)
    r_dest1.font.name = "Arial"

    p_dest2 = doc.add_paragraph()
    p_dest2.paragraph_format.space_before = Pt(0)
    p_dest2.paragraph_format.space_after = Pt(0)
    p_dest2.paragraph_format.line_spacing = 1.15

    r_dest2 = p_dest2.add_run("{{ DESTINATARIO_ENDERECO }}")
    r_dest2.font.size = Pt(11)
    r_dest2.font.name = "Arial"

    remover_quebras_manuais(doc)
    doc.save(str(destino))
    return destino


def gerar_envelope_combinado(
    destinatarios: list[tuple[str, str]],
    destino: str | Path,
) -> Path:
    """Generate a single envelope document with one page per recipient.

    Each page follows the DL envelope format (22 cm × 11 cm, landscape)
    and contains only the recipient name and address — no honorific prefix.

    Args:
        destinatarios: List of ``(nome, endereco)`` tuples where *nome* is the
            recipient name and *endereco* is the multi-line address string.
        destino: Target file path for the combined envelope document.

    Returns:
        Path of the file created.
    """
    from docx import Document  # type: ignore[import-untyped]  # noqa: PLC0415
    from docx.oxml import OxmlElement  # type: ignore[import-untyped]  # noqa: PLC0415
    from docx.oxml.ns import qn  # type: ignore[import-untyped]  # noqa: PLC0415
    from docx.shared import Cm, Pt  # type: ignore[import-untyped]  # noqa: PLC0415

    doc = Document()

    # Set page size to DL envelope (22 cm x 11 cm in landscape orientation)
    section = doc.sections[0]
    section.page_width = Cm(22)
    section.page_height = Cm(11)

    # Set margins
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    for i, (nome, endereco) in enumerate(destinatarios):
        # Add name paragraph
        p_nome = doc.add_paragraph()
        p_nome.paragraph_format.space_before = Pt(36)
        p_nome.paragraph_format.space_after = Pt(0)
        p_nome.paragraph_format.line_spacing = 1.15

        r_nome = p_nome.add_run(nome)
        r_nome.font.size = Pt(11)
        r_nome.font.name = "Arial"

        # Add address lines as separate paragraphs
        for linha in endereco.split("\n"):
            linha = linha.strip()
            if not linha:
                continue
            p_end = doc.add_paragraph()
            p_end.paragraph_format.space_before = Pt(0)
            p_end.paragraph_format.space_after = Pt(0)
            p_end.paragraph_format.line_spacing = 1.15

            r_end = p_end.add_run(linha)
            r_end.font.size = Pt(11)
            r_end.font.name = "Arial"

        # Add page break after each recipient except the last
        if i < len(destinatarios) - 1:
            p_break = doc.add_paragraph()
            p_break.paragraph_format.space_before = Pt(0)
            p_break.paragraph_format.space_after = Pt(0)
            run_elem = p_break._p
            br = OxmlElement("w:r")
            br_elem = OxmlElement("w:br")
            br_elem.set(qn("w:type"), "page")
            br.append(br_elem)
            run_elem.append(br)

    remover_quebras_manuais(doc)

    destino = Path(destino)
    destino.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(destino))
    return destino


def remover_quebras_manuais(doc: Any) -> None:
    """Substitui todas as quebras de linha manuais (Shift + Enter / <w:br/>) por parágrafos autônomos (<w:p>).

    Inspeciona todos os parágrafos do documento no corpo principal, cabeçalhos e rodapés.
    Cada parágrafo que contenha elementos <w:br/> ou quebras de linha em suas corridas
    é desmembrado em parágrafos separados (<w:p>), clonando as propriedades do parágrafo
    (<w:pPr>) e das corridas (<w:rPr>), garantindo a eliminação total de Shift + Enter.
    """
    from copy import deepcopy  # noqa: PLC0415
    from docx.oxml import OxmlElement  # type: ignore[import-untyped]  # noqa: PLC0415
    from docx.oxml.ns import qn  # type: ignore[import-untyped]  # noqa: PLC0415

    containers = [doc]
    for s in list(getattr(doc, "sections", [])):
        if getattr(s, "header", None):
            containers.append(s.header)
        if getattr(s, "footer", None):
            containers.append(s.footer)

    for container in containers:
        paragraphs = list(getattr(container, "paragraphs", []))
        for p in paragraphs:
            p_elem = p._p
            brs = p_elem.xpath(".//w:br")
            has_newline = False
            if not brs:
                for t_elem in p_elem.xpath(".//w:t"):
                    if t_elem.text and "\n" in t_elem.text:
                        has_newline = True
                        break

            if not brs and not has_newline:
                continue

            parent = p_elem.getparent()
            if parent is None:
                continue
            p_idx = parent.index(p_elem)
            pPr = p_elem.find(qn("w:pPr"))

            current_p = OxmlElement("w:p")
            if pPr is not None:
                current_p.append(deepcopy(pPr))

            new_paragraphs = [current_p]

            for child in list(p_elem):
                if child.tag.endswith("pPr"):
                    continue
                if child.tag.endswith("r"):
                    current_r = OxmlElement("w:r")
                    rPr = child.find(qn("w:rPr"))
                    if rPr is not None:
                        current_r.append(deepcopy(rPr))

                    for r_child in list(child):
                        if r_child.tag.endswith("rPr"):
                            continue
                        if r_child.tag.endswith("br"):
                            if len(current_r) > (1 if rPr is not None else 0):
                                current_p.append(current_r)

                            current_p = OxmlElement("w:p")
                            if pPr is not None:
                                current_p.append(deepcopy(pPr))
                            new_paragraphs.append(current_p)

                            current_r = OxmlElement("w:r")
                            if rPr is not None:
                                current_r.append(deepcopy(rPr))
                        elif r_child.tag.endswith("t"):
                            text = r_child.text or ""
                            if "\n" in text:
                                parts = text.split("\n")
                                for i, part in enumerate(parts):
                                    if i > 0:
                                        if len(current_r) > (1 if rPr is not None else 0):
                                            current_p.append(current_r)
                                        current_p = OxmlElement("w:p")
                                        if pPr is not None:
                                            current_p.append(deepcopy(pPr))
                                        new_paragraphs.append(current_p)
                                        current_r = OxmlElement("w:r")
                                        if rPr is not None:
                                            current_r.append(deepcopy(rPr))
                                    if part:
                                        t_elem = OxmlElement("w:t")
                                        t_elem.text = part
                                        if part.startswith(" ") or part.endswith(" "):
                                            t_elem.set(qn("xml:space"), "preserve")
                                        current_r.append(t_elem)
                            else:
                                current_r.append(deepcopy(r_child))
                        else:
                            current_r.append(deepcopy(r_child))

                    if len(current_r) > (1 if rPr is not None else 0):
                        current_p.append(current_r)
                else:
                    current_p.append(deepcopy(child))

            for offset, new_p in enumerate(new_paragraphs):
                parent.insert(p_idx + offset, new_p)
            parent.remove(p_elem)


def ajustar_posicao_rodape(caminho_doc: str, word_app: Any = None) -> None:
    """Ajusta o espaçamento antes do bloco de destinatário (rodapé do ofício)
    para que ele termine exatamente na penúltima linha da primeira página.
    
    Usa o Word COM Automation (win32com). Se `word_app` não for fornecido,
    cria uma instância temporária do Word.
    """
    import os
    from pathlib import Path
    
    path = Path(caminho_doc)
    if not path.exists():
        return
        
    try:
        import win32com.client
    except ImportError:
        return
        
    temp_word = False
    doc = None
    try:
        if word_app is None:
            try:
                word_app = win32com.client.DispatchEx("Word.Application")
                word_app.Visible = False
                temp_word = True
            except Exception:
                return
                
        doc = word_app.Documents.Open(str(path.resolve()))
        
        def get_page(p: Any) -> int:
            try:
                return int(p.Range.Information(3))  # 3 represents wdActiveEndPageNumber
            except Exception:
                return 1
                
        n_paragraphs = doc.paragraphs.Count
        if n_paragraphs < 1:
            return
            
        # Identificar o bloco de destinatário (parágrafos não vazios no final)
        p_idx = n_paragraphs
        while p_idx >= 1:
            try:
                text = doc.paragraphs(p_idx).Range.Text.strip()
            except Exception:
                text = ""
            if not text:
                break
            p_idx -= 1
            
        idx_recip_start = p_idx + 1
        
        if idx_recip_start > n_paragraphs:
            return
            
        insert_idx = idx_recip_start - 1
        if insert_idx < 1:
            return
            
        p_last = doc.paragraphs(n_paragraphs)
        
        # 1. Se estiver na página 2+, remove parágrafos vazios anteriores para tentar caber na página 1
        max_deletes = 100
        while get_page(p_last) > 1 and insert_idx > 1 and max_deletes > 0:
            prev_p = doc.paragraphs(insert_idx)
            try:
                p_text = prev_p.Range.Text.strip()
            except Exception:
                p_text = "error"
            if not p_text:
                prev_p.Range.Delete()
                n_paragraphs = doc.paragraphs.Count
                p_last = doc.paragraphs(n_paragraphs)
                insert_idx -= 1
            else:
                break
            max_deletes -= 1
            
        # 2. Se estiver na página 1, insere parágrafos vazios até que transpasse para a página 2
        if get_page(p_last) == 1:
            spilled = False
            added_count = 0
            max_inserts = 100
            
            while get_page(p_last) == 1 and max_inserts > 0:
                p_recip = doc.paragraphs(idx_recip_start + added_count)
                p_recip.Range.InsertParagraphBefore()
                added_count += 1
                
                n_paragraphs = doc.paragraphs.Count
                p_last = doc.paragraphs(n_paragraphs)
                
                if get_page(p_last) > 1:
                    spilled = True
                    break
                max_inserts -= 1
                
            # 3. Uma vez que transpassou, remove exatamente 2 parágrafos vazios
            # (1 para voltar à página 1 na última linha, e mais 1 para ficar na penúltima linha)
            if spilled:
                for _ in range(2):
                    n_paragraphs = doc.paragraphs.Count
                    empty_idx = idx_recip_start + added_count - 1
                    if empty_idx >= 1:
                        p_to_del = doc.paragraphs(empty_idx)
                        try:
                            p_to_del_text = p_to_del.Range.Text.strip()
                        except Exception:
                            p_to_del_text = "error"
                        if not p_to_del_text:
                            p_to_del.Range.Delete()
                            added_count -= 1
                            
        doc.Save()
    except Exception:
        pass
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if temp_word and word_app is not None:
            try:
                word_app.Quit()
            except Exception:
                pass


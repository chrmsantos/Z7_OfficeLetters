"""Conferência e correção automática dos ofícios e planilha gerados.

Implementa a Fase 6 do fluxo de processamento: após a geração de todos os
ofícios (.docx), este módulo verifica a consistência e a correção linguística
de cada documento antes de a planilha Excel ser gravada em disco.

Verificações realizadas por ofício:
- **Consistência de dados**: ctx ↔ dados extraídos pela IA (números de moção,
  tipos, falecido, tipo_propositura).
- **Concordância linguística**: consistência interna entre vocativo,
  pronome_corpo e tratamento_rodapé; concordância de número nas designações,
  artigos e particípios.
- **Consistência da planilha**: cada linha Excel é conferida contra o ctx do
  ofício correspondente.

Erros encontrados são corrigidos automaticamente por re-renderização do
template .docx; erros irrecuperáveis são registrados no log e o fluxo segue.

Public exports:
    RegistroOficio: snapshot de um ofício gerado, preenchido pelo processor.
    ResultadoVerificacao: resultado da verificação de um único ofício.
    RelatorioConferencia: relatório agregado de toda a rodada.
    conferir_trabalho: orquestrador principal — recebe os registros e executa tudo.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from typing import Any

from z7_officeletters.core.documents import formatar_lista_pt as _formatar_lista_pt
from z7_officeletters.core.documents import frases_propositura as _frases_propositura
from z7_officeletters.core.documents import _titlecase_nome

__all__ = [
    "RegistroOficio",
    "ResultadoVerificacao",
    "RelatorioConferencia",
    "conferir_trabalho",
]

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------


@dataclass
class RegistroOficio:
    """Snapshot de um ofício gerado, mantido para verificação na Fase 6.

    Attributes:
        caminho: Caminho absoluto para o arquivo .docx gerado.
        nome_arquivo: Somente o nome do arquivo (sem diretório).
        ctx: Dicionário de contexto usado para renderizar o template (mutável —
            é atualizado in-place quando correções são aplicadas).
        dados_grupo: Lista dos dicts de dados extraídos pela IA para o grupo
            de proposituras agrupadas neste ofício.
        dest_raw: Dict bruto do destinatário extraído pela IA (primeiro do grupo).
        info: Dict ``DestinatarioProcessado`` (já com sobrescritas do DB de
            endereços, quando aplicável).
        n_props: Número de proposituras agrupadas neste ofício.
        tipo_propositura: ``"mocao"`` ou ``"requerimento_pesar"``.
        template_path: Caminho absoluto do template .docx utilizado.
        linha_planilha_idx: Índice da linha correspondente em ``dados_planilha``.
    """

    caminho: str
    nome_arquivo: str
    ctx: dict[str, str]
    dados_grupo: list[dict[str, Any]]
    dest_raw: dict[str, Any]
    info: dict[str, str]
    n_props: int
    tipo_propositura: str
    template_path: str
    linha_planilha_idx: int


@dataclass
class ResultadoVerificacao:
    """Resultado da verificação de um único ofício.

    Attributes:
        arquivo: Nome do arquivo verificado.
        erros_dados: Erros de consistência de dados (números, tipos, falecido).
        erros_linguisticos: Erros de concordância linguística.
        erros_planilha: Erros na linha correspondente da planilha Excel.
        erros_tags: Marcadores de template não resolvidos no arquivo gerado.
        corrigido: True quando todos os erros foram corrigidos com sucesso.
        incorrigivel: True quando houve erros mas a correção falhou.
    """

    arquivo: str
    erros_dados: list[str] = field(default_factory=list)
    erros_linguisticos: list[str] = field(default_factory=list)
    erros_planilha: list[str] = field(default_factory=list)
    erros_tags: list[str] = field(default_factory=list)
    corrigido: bool = False
    incorrigivel: bool = False

    @property
    def tem_erros(self) -> bool:
        """True se houver qualquer erro encontrado."""
        return bool(
            self.erros_dados
            or self.erros_linguisticos
            or self.erros_planilha
            or self.erros_tags
        )

    @property
    def todos_erros(self) -> list[str]:
        """Lista plana de todos os erros encontrados."""
        return (
            self.erros_dados
            + self.erros_linguisticos
            + self.erros_planilha
            + self.erros_tags
        )


def verificar_tags_pendentes(caminho: str) -> list[str]:
    """Verifica se o arquivo final renderizado contém marcadores pendentes de template (como {{ ou }})."""
    from docx import Document  # noqa: PLC0415
    import os  # noqa: PLC0415

    erros = []
    if not os.path.exists(caminho):
        erros.append(f"Arquivo não encontrado para conferência de tags: {caminho}")
        return erros

    try:
        doc = Document(caminho)
        # 1. Verificar parágrafos do corpo
        for p_idx, p in enumerate(doc.paragraphs, start=1):
            if "{{" in p.text or "}}" in p.text or "{%" in p.text or "%}" in p.text:
                erros.append(f"Marcador de template não renderizado no parágrafo {p_idx}: '{p.text}'")

        # 2. Verificar tabelas
        for t_idx, table in enumerate(doc.tables, start=1):
            for r_idx, row in enumerate(table.rows, start=1):
                for c_idx, cell in enumerate(row.cells, start=1):
                    if "{{" in cell.text or "}}" in cell.text or "{%" in cell.text or "%}" in cell.text:
                        erros.append(
                            f"Marcador de template não renderizado na tabela {t_idx}, linha {r_idx}, célula {c_idx}: '{cell.text}'"
                        )

        # 3. Verificar cabeçalhos e rodapés de cada seção (incluindo primeira página e páginas pares)
        for s_idx, section in enumerate(doc.sections, start=1):
            headers = []
            footers = []
            
            if hasattr(section, "header") and section.header is not None:
                headers.append(("cabeçalho padrão", section.header))
            if hasattr(section, "first_page_header") and section.first_page_header is not None:
                headers.append(("cabeçalho da primeira página", section.first_page_header))
            if hasattr(section, "even_page_header") and section.even_page_header is not None:
                headers.append(("cabeçalho de página par", section.even_page_header))
                
            if hasattr(section, "footer") and section.footer is not None:
                footers.append(("rodapé padrão", section.footer))
            if hasattr(section, "first_page_footer") and section.first_page_footer is not None:
                footers.append(("rodapé da primeira página", section.first_page_footer))
            if hasattr(section, "even_page_footer") and section.even_page_footer is not None:
                footers.append(("rodapé de página par", section.even_page_footer))

            for tipo_lbl, h in headers:
                for p_idx, p in enumerate(h.paragraphs, start=1):
                    if "{{" in p.text or "}}" in p.text or "{%" in p.text or "%}" in p.text:
                        erros.append(
                            f"Marcador de template não renderizado no {tipo_lbl} da seção {s_idx}, parágrafo {p_idx}: '{p.text}'"
                        )

            for tipo_lbl, f in footers:
                for p_idx, p in enumerate(f.paragraphs, start=1):
                    if "{{" in p.text or "}}" in p.text or "{%" in p.text or "%}" in p.text:
                        erros.append(
                            f"Marcador de template não renderizado no {tipo_lbl} da seção {s_idx}, parágrafo {p_idx}: '{p.text}'"
                        )
    except Exception as e:
        erros.append(f"Erro ao analisar marcadores do arquivo '{caminho}': {e}")

    return erros


@dataclass
class RelatorioConferencia:
    """Relatório agregado de toda a rodada de conferência.

    Attributes:
        total_verificados: Total de ofícios processados pela verificação.
        total_com_erros: Ofícios em que pelo menos um erro foi encontrado.
        total_corrigidos: Ofícios com erros que foram corrigidos com sucesso.
        total_incorrigiveis: Ofícios com erros que não puderam ser corrigidos.
        resultados: Lista de :class:`ResultadoVerificacao` por ofício.
    """

    total_verificados: int = 0
    total_com_erros: int = 0
    total_corrigidos: int = 0
    total_incorrigiveis: int = 0
    resultados: list[ResultadoVerificacao] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Verification functions
# ---------------------------------------------------------------------------


def verificar_consistencia_dados(registro: RegistroOficio) -> list[str]:
    """Verifica se os valores do ctx correspondem aos dados brutos da IA.

    Verifica:
    - ``num_mocao`` no ctx é o merge correto dos números em ``dados_grupo``.
    - ``tipo_mocao`` no ctx é o merge correto dos tipos em ``dados_grupo``.
    - ``falecido`` no ctx corresponde ao merge dos falecidos (requerimento_pesar).
    - ``tipo_propositura`` no ctx bate com o tipo do registro.

    Returns:
        Lista de descrições de erros encontrados (vazia = sem erros).
    """
    erros: list[str] = []
    ctx = registro.ctx
    dados = registro.dados_grupo

    # Re-compute expected merged values from source AI data
    nums_lista = [d["numero_mocao"] for d in dados if d.get("numero_mocao")]
    nums_esperados = _formatar_lista_pt(nums_lista) if nums_lista else ""

    tipos_lista = [str(d.get("tipo_mocao", "")) for d in dados if d.get("tipo_mocao")]
    tipo_esperado = _formatar_lista_pt(tipos_lista) if tipos_lista else ""

    falecidos_lista = [str(d.get("falecido", "")) for d in dados if d.get("falecido")]
    falecido_esperado = _formatar_lista_pt(falecidos_lista) if falecidos_lista else ""

    # Check num_mocao
    if nums_esperados and ctx.get("num_mocao") != nums_esperados:
        erros.append(
            f"num_mocao: ctx='{ctx.get('num_mocao')}' — esperado='{nums_esperados}'"
        )

    # Check tipo_mocao (only for moções)
    if registro.tipo_propositura != "requerimento_pesar":
        if ctx.get("tipo_mocao", "") != tipo_esperado:
            erros.append(
                f"tipo_mocao: ctx='{ctx.get('tipo_mocao')}' — esperado='{tipo_esperado}'"
            )

    # Check falecido (only for requerimentos de pesar)
    if registro.tipo_propositura == "requerimento_pesar":
        if ctx.get("falecido", "") != falecido_esperado:
            erros.append(
                f"falecido: ctx='{ctx.get('falecido')}' — esperado='{falecido_esperado}'"
            )

    # Check tipo_propositura field itself
    if ctx.get("tipo_propositura", "") != registro.tipo_propositura:
        erros.append(
            f"tipo_propositura: ctx='{ctx.get('tipo_propositura')}' "
            f"— esperado='{registro.tipo_propositura}'"
        )

    # Check for placeholder/empty/AI-hallucinated values in all ctx keys
    placeholders = {"none", "null", "todo", "indef", "indefinido", "unknown"}
    for chave, valor in ctx.items():
        if valor is None:
            erros.append(f"Campo '{chave}' no contexto está nulo (None)")
        else:
            val_strip = str(valor).strip().lower()
            if val_strip in placeholders or not val_strip:
                # Except empty fields that are expected to be optionally empty
                chave_lower = chave.lower()
                if registro.tipo_propositura == "requerimento_pesar":
                    if chave_lower in ("tipo_mocao", "destinatario_endereco"):
                        continue
                else:
                    if chave_lower in ("falecido", "destinatario_endereco"):
                        continue
                erros.append(f"Campo '{chave}' no contexto possui valor inválido/incompleto: '{valor}'")

    return erros


def verificar_concordancia_linguistica(registro: RegistroOficio) -> list[str]:
    """Verifica concordância de gênero, número e pronomes no ctx.

    Regras verificadas:
    - Consistência interna: vocativo ↔ pronome_corpo ↔ tratamento_rodapé.
    - Requerimento de pesar: formas fixas obrigatórias.
    - Concordância de número: designacao_propositura / copia_art / aprovada_s
      batem com ``n_props``.
    - Prefeito (``is_prefeito``): formas fixas e destinatário correto.
    - Instituição (PJ/Coletivo): formas plurais.

    Returns:
        Lista de descrições de erros encontrados (vazia = sem erros).
    """
    erros: list[str] = []
    ctx = registro.ctx
    dest = registro.dest_raw
    n_props = registro.n_props
    tipo_prop = registro.tipo_propositura
    tipo_mocao = ctx.get("tipo_mocao", "")

    # ── Requerimento de pesar: formas fixas ──────────────────────────────────
    if tipo_prop == "requerimento_pesar":
        if ctx.get("vocativo") != "Ilustríssimos Senhores(as)":
            erros.append(
                f"Requerimento de pesar — vocativo: "
                f"ctx='{ctx.get('vocativo')}' esperado='Ilustríssimos Senhores(as)'"
            )
        if ctx.get("pronome_corpo") != "Vossas Senhorias":
            erros.append(
                f"Requerimento de pesar — pronome_corpo: "
                f"ctx='{ctx.get('pronome_corpo')}' esperado='Vossas Senhorias'"
            )
        _verificar_numero(erros, ctx, tipo_prop, tipo_mocao, n_props)
        return erros

    # ── Internal vocativo ↔ pronome_corpo consistency ────────────────────────
    vocativo = ctx.get("vocativo", "")
    pronome = ctx.get("pronome_corpo", "")
    tratamento = ctx.get("tratamento_rodape", "")
    _verificar_consistencia_pronomes(erros, vocativo, pronome, tratamento)

    # ── Recipient-specific checks ────────────────────────────────────────────
    tipo_dest = dest.get("tipo", "PF")
    is_inst = tipo_dest in ("PJ", "Coletivo") or bool(dest.get("is_instituicao"))
    is_prefeito = bool(dest.get("is_prefeito")) or "prefeito" in (dest.get("nome") or "").lower()

    if is_prefeito:
        _verificar_prefeito(erros, ctx)
    elif is_inst:
        _verificar_instituicao(erros, ctx, dest)
    else:
        genero = dest.get("genero") or "M"
        nivel = dest.get("nivel_protocolo") or "VS"
        _verificar_pf(erros, ctx, dest, genero, nivel)

    # ── Number agreement ─────────────────────────────────────────────────────
    _verificar_numero(erros, ctx, tipo_prop, tipo_mocao, n_props)

    return erros


def _verificar_consistencia_pronomes(
    erros: list[str], vocativo: str, pronome: str, tratamento: str
) -> None:
    """Verifica a consistência interna entre vocativo, pronome e tratamento."""
    voc_lower = vocativo.lower()

    if "excelentíssim" in voc_lower:
        if pronome != "Vossa Excelência":
            erros.append(
                f"Inconsistência: vocativo '{vocativo}' requer pronome "
                f"'Vossa Excelência' — encontrado '{pronome}'"
            )
    elif "ilustríssimos" in voc_lower or "ilustríssimas" in voc_lower:
        if pronome != "Vossas Senhorias":
            erros.append(
                f"Inconsistência: vocativo plural '{vocativo}' requer "
                f"'Vossas Senhorias' — encontrado '{pronome}'"
            )
    elif "reverendíssim" in voc_lower or "reverendissim" in voc_lower:
        if pronome != "Vossa Reverendíssima":
            erros.append(
                f"Inconsistência: vocativo '{vocativo}' requer pronome "
                f"'Vossa Reverendíssima' — encontrado '{pronome}'"
            )
    elif "ilustríssim" in voc_lower:
        if pronome != "Vossa Senhoria":
            erros.append(
                f"Inconsistência: vocativo '{vocativo}' requer "
                f"'Vossa Senhoria' — encontrado '{pronome}'"
            )

    # Gender consistency between vocativo and tratamento_rodape
    voc_feminino = "senhora" in voc_lower and "senhori" not in voc_lower
    voc_masculino = (
        "senhor" in voc_lower
        and "senhora" not in voc_lower
        and "senhori" not in voc_lower
        and "senhorias" not in voc_lower
    )
    trat_lower = tratamento.lower()
    if voc_feminino and "senhor" in trat_lower and "senhora" not in trat_lower:
        erros.append(
            f"Inconsistência: vocativo feminino '{vocativo}' mas "
            f"tratamento_rodapé com forma masculina '{tratamento}'"
        )
    if voc_masculino and "senhora" in trat_lower and "senhori" not in trat_lower:
        erros.append(
            f"Inconsistência: vocativo masculino '{vocativo}' mas "
            f"tratamento_rodapé com forma feminina '{tratamento}'"
        )


def _verificar_prefeito(erros: list[str], ctx: dict[str, str]) -> None:
    """Verifica que o ctx usa as formas fixas obrigatórias para o Prefeito."""
    from z7_officeletters.core.config import PREFEITO  # noqa: PLC0415

    if ctx.get("vocativo") != "Excelentíssimo Senhor Prefeito":
        erros.append(
            f"Prefeito — vocativo: ctx='{ctx.get('vocativo')}' "
            f"esperado='Excelentíssimo Senhor Prefeito'"
        )
    if ctx.get("pronome_corpo") != "Vossa Excelência":
        erros.append(
            f"Prefeito — pronome_corpo: ctx='{ctx.get('pronome_corpo')}' "
            f"esperado='Vossa Excelência'"
        )
    nome_esperado = PREFEITO.get("nome", "").upper()
    if nome_esperado and ctx.get("destinatario_nome") != nome_esperado:
        erros.append(
            f"Prefeito — destinatario_nome: ctx='{ctx.get('destinatario_nome')}' "
            f"esperado='{nome_esperado}'"
        )


def _verificar_instituicao(erros: list[str], ctx: dict[str, str], dest: dict[str, Any]) -> None:
    """Verifica que o ctx usa as formas corretas para PJ/Coletivo.

    Quando a instituição possui representante nomeado, as formas devem ser
    singulares (endereçadas à pessoa); sem representante, as formas plurais
    genéricas são esperadas.
    """
    pronome = ctx.get("pronome_corpo", "")
    vocativo = ctx.get("vocativo", "")
    representante = (dest.get("representante") or "").strip()

    if representante:
        from z7_officeletters.core.recipients import _is_clergy
        funcao_rep = dest.get("funcao_representante") or ""
        if _is_clergy(funcao_rep, representante):
            if pronome != "Vossa Reverendíssima":
                erros.append(
                    f"Instituição com representante do clero — pronome_corpo: ctx='{pronome}' "
                    f"esperado='Vossa Reverendíssima' (singular, endereçado a '{representante}')"
                )
            voc_lower = vocativo.lower()
            if "reverendíssima senhora" not in voc_lower and "reverendíssimo senhor" not in voc_lower:
                erros.append(
                    f"Instituição com representante do clero — vocativo: ctx='{vocativo}' "
                    f"deve ser singular ('Reverendíssima Senhora' ou 'Reverendíssimo Senhor')"
                )
        else:
            # Institution with a named representative → singular forms expected.
            if pronome != "Vossa Senhoria":
                erros.append(
                    f"Instituição com representante — pronome_corpo: ctx='{pronome}' "
                    f"esperado='Vossa Senhoria' (singular, endereçado a '{representante}')"
                )
            voc_lower = vocativo.lower()
            if "ilustríssima senhora" not in voc_lower and "ilustríssimo senhor" not in voc_lower:
                erros.append(
                    f"Instituição com representante — vocativo: ctx='{vocativo}' "
                    f"deve ser singular ('Ilustríssima Senhora' ou 'Ilustríssimo Senhor')"
                )
    else:
        # No representative → generic plural forms expected.
        if pronome != "Vossas Senhorias":
            erros.append(
                f"Instituição — pronome_corpo: ctx='{pronome}' esperado='Vossas Senhorias'"
            )
        voc_lower = vocativo.lower()
        if "senhores" not in voc_lower and "senhoras" not in voc_lower:
            erros.append(
                f"Instituição — vocativo: ctx='{vocativo}' não está no plural "
                f"(esperado 'Ilustríssimos Senhores' ou 'Ilustríssimas Senhoras')"
            )


def _verificar_pf(
    erros: list[str], ctx: dict[str, str], dest: dict[str, Any], genero: str, nivel: str
) -> None:
    """Verifica formas de gênero e nível de protocolo para pessoa física."""
    vocativo = ctx.get("vocativo", "")
    tratamento = ctx.get("tratamento_rodape", "")
    voc_lower = vocativo.lower()
    trat_lower = tratamento.lower()

    import re

    if genero == "F":
        if "senhor" in voc_lower and "senhora" not in voc_lower:
            erros.append(
                f"Gênero F — vocativo com forma masculina: '{vocativo}'"
            )
        # Strict check for male terms in female context
        for term in ["senhor", "senhores", "ilustríssimo", "ilustríssimos", "excelentíssimo", "excelentíssimos", "reverendíssimo", "ao"]:
            pattern = rf"\b{term}\b"
            if re.search(pattern, voc_lower) or re.search(pattern, trat_lower):
                erros.append(f"Gênero F — termo masculino '{term}' encontrado no vocativo/tratamento")
        if nivel == "VS":
            from z7_officeletters.core.recipients import _is_clergy
            funcao_prof = dest.get("funcao_profissao") or ""
            nome = dest.get("nome") or ""
            if _is_clergy(funcao_prof, nome):
                if "reverendíssima" not in voc_lower:
                    erros.append(
                        f"Gênero F, clero, nível VS — vocativo: '{vocativo}' "
                        f"deveria conter 'Reverendíssima'"
                    )
                if not tratamento.startswith("À Reverendíssima"):
                    erros.append(
                        f"Gênero F, clero, nível VS — tratamento_rodapé: '{tratamento}' "
                        f"deveria começar com 'À Reverendíssima'"
                    )
            else:
                if "ilustríssima" not in voc_lower:
                    erros.append(
                        f"Gênero F, nível VS — vocativo: '{vocativo}' "
                        f"deveria conter 'Ilustríssima'"
                    )
                if not tratamento.startswith("À Ilustríssima"):
                    erros.append(
                        f"Gênero F, nível VS — tratamento_rodapé: '{tratamento}' "
                        f"deveria começar com 'À Ilustríssima'"
                    )
        elif nivel in ("VE", "VE_M"):
            if "excelentíssima" not in voc_lower:
                erros.append(
                    f"Gênero F, nível {nivel} — vocativo: '{vocativo}' "
                    f"deveria conter 'Excelentíssima'"
                )
    else:  # genero == "M"
        # Strict check for female terms in male context
        for term in ["senhora", "senhoras", "ilustríssima", "ilustríssimas", "excelentíssima", "excelentíssimas", "reverendíssima", "à"]:
            pattern = rf"\b{term}\b"
            if re.search(pattern, voc_lower) or re.search(pattern, trat_lower):
                erros.append(f"Gênero M — termo feminino '{term}' encontrado no vocativo/tratamento")
        if nivel == "VS":
            from z7_officeletters.core.recipients import _is_clergy
            funcao_prof = dest.get("funcao_profissao") or ""
            nome = dest.get("nome") or ""
            if _is_clergy(funcao_prof, nome):
                if "reverendíssimo" not in voc_lower:
                    erros.append(
                        f"Gênero M, clero, nível VS — vocativo: '{vocativo}' "
                        f"deveria conter 'Reverendíssimo'"
                    )
                if not tratamento.startswith("Ao Reverendíssimo"):
                    erros.append(
                        f"Gênero M, clero, nível VS — tratamento_rodapé: '{tratamento}' "
                        f"deveria começar com 'Ao Reverendíssimo'"
                    )
            else:
                if "ilustríssimo" not in voc_lower:
                    erros.append(
                        f"Gênero M, nível VS — vocativo: '{vocativo}' "
                        f"deveria conter 'Ilustríssimo'"
                    )
                if not tratamento.startswith("Ao Ilustríssimo"):
                    erros.append(
                        f"Gênero M, nível VS — tratamento_rodapé: '{tratamento}' "
                        f"deveria começar com 'Ao Ilustríssimo'"
                    )
        elif nivel in ("VE", "VE_M"):
            if "excelentíssimo" not in voc_lower:
                erros.append(
                    f"Gênero M, nível {nivel} — vocativo: '{vocativo}' "
                    f"deveria conter 'Excelentíssimo'"
                )

    # Protocol-level crase check
    if nivel == "VE":
        if "À Sua Excelência" in tratamento:
            erros.append(
                f"Nível VE (federal/estadual) — tratamento_rodapé: '{tratamento}' "
                f"não deve ter crase; esperado 'A Sua Excelência'"
            )
    elif nivel == "VE_M":
        if tratamento.startswith("A Sua Excelência") and not tratamento.startswith("À"):
            erros.append(
                f"Nível VE_M (municipal) — tratamento_rodapé: '{tratamento}' "
                f"deve ter crase; esperado 'À Sua Excelência'"
            )


def _verificar_numero(
    erros: list[str],
    ctx: dict[str, str],
    tipo_prop: str,
    tipo_mocao: str,
    n_props: int,
) -> None:
    """Verifica concordância de número em designacao, copia_art e aprovada_s."""
    esp_desig, esp_art, esp_aprov = _frases_propositura(tipo_prop, tipo_mocao, n_props)

    ctx_desig = ctx.get("designacao_propositura", "")
    ctx_art = ctx.get("copia_art", "")
    ctx_aprov = ctx.get("aprovada_s", "")

    if ctx_desig and ctx_desig != esp_desig:
        erros.append(
            f"Designação (n={n_props}): ctx='{ctx_desig}' — esperado='{esp_desig}'"
        )
    if ctx_art and ctx_art != esp_art:
        erros.append(
            f"Artigo cópia (n={n_props}): ctx='{ctx_art}' — esperado='{esp_art}'"
        )
    if ctx_aprov and ctx_aprov != esp_aprov:
        erros.append(
            f"Particípio (n={n_props}): ctx='{ctx_aprov}' — esperado='{esp_aprov}'"
        )


def verificar_linha_planilha(
    linha: list[Any], registro: RegistroOficio
) -> list[str]:
    """Verifica se a linha da planilha é consistente com o ofício correspondente.

    Confere:
    - ``linha[0]`` (num_oficio) == ctx[``num_oficio``].
    - ``linha[2]`` (destinatário) == tratamento_rodapé + destinatario_nome.
    - ``linha[3]`` (assunto) contém o tipo e os números de moção corretos.
    - ``linha[5]`` (envio) == info[``envio``].

    Returns:
        Lista de descrições de erros encontrados (vazia = sem erros).
    """
    erros: list[str] = []
    ctx = registro.ctx
    info = registro.info
    n_props = registro.n_props
    tipo_prop = registro.tipo_propositura
    num_mocao = ctx.get("num_mocao", "")
    tipo_mocao = ctx.get("tipo_mocao", "")

    if len(linha) < 6:
        erros.append(f"Linha da planilha com apenas {len(linha)} coluna(s) — esperado ≥ 6")
        return erros

    num_planilha = str(linha[0])
    dest_planilha = str(linha[2])
    assunto_planilha = str(linha[3])
    envio_planilha = str(linha[5])

    # num_oficio
    if num_planilha != ctx.get("num_oficio", ""):
        erros.append(
            f"Planilha num_oficio: '{num_planilha}' ≠ ctx '{ctx.get('num_oficio')}'"
        )

    # destinatário
    dest_nome_title = _titlecase_nome(info.get('destinatario_nome', ''))
    esperado_dest = (
        f"{info.get('tratamento_rodape', '')} {dest_nome_title}".strip()
    )
    if dest_planilha != esperado_dest:
        erros.append(
            f"Planilha destinatário: '{dest_planilha}' ≠ esperado '{esperado_dest}'"
        )

    # assunto: deve conter o número da moção
    if num_mocao and num_mocao not in assunto_planilha:
        erros.append(
            f"Planilha assunto: '{assunto_planilha}' não contém número '{num_mocao}'"
        )

    # assunto: deve conter o tipo (para moções)
    if tipo_prop != "requerimento_pesar" and tipo_mocao and tipo_mocao not in assunto_planilha:
        erros.append(
            f"Planilha assunto: '{assunto_planilha}' não contém tipo '{tipo_mocao}'"
        )

    # envio
    if envio_planilha != info.get("envio", ""):
        erros.append(
            f"Planilha envio: '{envio_planilha}' ≠ '{info.get('envio')}'"
        )

    return erros


# ---------------------------------------------------------------------------
# Correction
# ---------------------------------------------------------------------------


def corrigir_ctx(
    registro: RegistroOficio,
    erros_dados: list[str],
    erros_ling: list[str],
) -> dict[str, str]:
    """Deriva um ctx corrigido a partir dos dados de origem e do info processado.

    Aplica correções cirúrgicas: apenas campos com erros detectados são
    recalculados, preservando os demais valores originais (inclusive possíveis
    sobrescritas vindas do banco de endereços).

    Args:
        registro: Registro do ofício a corrigir.
        erros_dados: Lista de erros de consistência de dados (de
            :func:`verificar_consistencia_dados`).
        erros_ling: Lista de erros linguísticos (de
            :func:`verificar_concordancia_linguistica`).

    Returns:
        Novo dict ctx com as correções aplicadas (o ctx original não é mutado).
    """
    ctx = dict(registro.ctx)  # shallow copy — str values are immutable
    dados = registro.dados_grupo
    info = registro.info
    n_props = registro.n_props
    tipo_prop = registro.tipo_propositura

    # ── Fix data fields ───────────────────────────────────────────────────────
    erros_dados_lower = [e.lower() for e in erros_dados]

    if any("num_mocao" in e for e in erros_dados_lower):
        nums_lista = [d["numero_mocao"] for d in dados if d.get("numero_mocao")]
        if nums_lista:
            merged = _formatar_lista_pt(nums_lista)
            ctx["num_mocao"] = merged
            ctx["NUM_MOCAO"] = merged

    if any("tipo_mocao" in e for e in erros_dados_lower):
        tipos = [str(d.get("tipo_mocao", "")) for d in dados if d.get("tipo_mocao")]
        if tipos:
            merged = _formatar_lista_pt(tipos)
            ctx["tipo_mocao"] = merged
            ctx["TIPO_MOCAO"] = merged

    if any("falecido" in e for e in erros_dados_lower):
        falecidos = [str(d.get("falecido", "")) for d in dados if d.get("falecido")]
        if falecidos:
            merged = _formatar_lista_pt(falecidos)
            ctx["falecido"] = merged
            ctx["FALECIDO"] = merged

    if any("tipo_propositura" in e for e in erros_dados_lower):
        ctx["tipo_propositura"] = tipo_prop
        ctx["TIPO_PROPOSITURA"] = tipo_prop

    # ── Fix honorifics / protocol forms ──────────────────────────────────────
    # Use ``info`` (already contains DB overrides) to avoid regression.
    erros_ling_lower = [e.lower() for e in erros_ling]
    honorific_keywords = ("vocativo", "pronome", "gênero", "tratamento", "prefeito", "instituição")
    if any(kw in e for e in erros_ling_lower for kw in honorific_keywords):
        ctx["vocativo"] = info["vocativo"]
        ctx["VOCATIVO"] = info["vocativo"]
        ctx["pronome_corpo"] = info["pronome_corpo"]
        ctx["PRONOME_CORPO"] = info["pronome_corpo"]
        ctx["tratamento_rodape"] = info["tratamento_rodape"]
        ctx["TRATAMENTO_RODAPE"] = info["tratamento_rodape"]
        ctx["destinatario_nome"] = info["destinatario_nome"]
        ctx["DESTINATARIO_NOME"] = info["destinatario_nome"]

    # ── Fix requerimento_pesar fixed overrides ────────────────────────────────
    if tipo_prop == "requerimento_pesar":
        ctx["vocativo"] = "Ilustríssimos Senhores(as)"
        ctx["VOCATIVO"] = "Ilustríssimos Senhores(as)"
        ctx["pronome_corpo"] = "Vossas Senhorias"
        ctx["PRONOME_CORPO"] = "Vossas Senhorias"

    # ── Fix number agreement ──────────────────────────────────────────────────
    number_keywords = ("designação", "artigo", "particípio", "copia", "aprovada", "designacao")
    if any(kw in e for e in erros_ling_lower for kw in number_keywords):
        tipo_mocao = ctx.get("tipo_mocao", "")
        desig, art, aprov = _frases_propositura(tipo_prop, tipo_mocao, n_props)
        ctx["designacao_propositura"] = desig
        ctx["DESIGNACAO_PROPOSITURA"] = desig
        ctx["copia_art"] = art
        ctx["COPIA_ART"] = art
        ctx["aprovada_s"] = aprov
        ctx["APROVADA_S"] = aprov

    return ctx


def _corrigir_linha_planilha(
    linha: list[Any], registro: RegistroOficio
) -> list[Any]:
    """Retorna uma linha de planilha corrigida para o registro fornecido."""
    ctx = registro.ctx
    info = registro.info
    n_props = registro.n_props
    tipo_prop = registro.tipo_propositura

    # Pad line to at least 6 elements to prevent IndexError
    linha_corr = list(linha)
    if len(linha_corr) < 6:
        linha_corr.extend([""] * (6 - len(linha_corr)))

    # Extrair o ano da data (campo linha_corr[1], formato YYYY-MM-DD ou DD/MM/AAAA)
    data_str = str(linha_corr[1])
    if "-" in data_str:
        year = data_str[:4] if len(data_str) >= 4 else ""
    elif "/" in data_str:
        year = data_str[-4:] if len(data_str) >= 4 else ""
    else:
        year = data_str[:4] if len(data_str) >= 4 else ""

    num_mocao = ctx.get("num_mocao", "")
    tipo_mocao = ctx.get("tipo_mocao", "")

    if tipo_prop == "requerimento_pesar":
        plural_s = "s" if n_props > 1 else ""
        assunto = f"Encaminha Requerimento{plural_s} de Pesar nº {num_mocao}/{year}"
    else:
        plural_oes = "ções" if n_props > 1 else "ção"
        assunto = f"Encaminha Mo{plural_oes} de {tipo_mocao} nº {num_mocao}/{year}"

    dest_nome_title = _titlecase_nome(info.get('destinatario_nome', ''))
    destinatario = (
        f"{info.get('tratamento_rodape', '')} {dest_nome_title}".strip()
    )

    linha_corr[0] = ctx.get("num_oficio", linha_corr[0])
    linha_corr[2] = destinatario
    linha_corr[3] = assunto
    linha_corr[5] = info.get("envio", linha_corr[5])
    return linha_corr


# ---------------------------------------------------------------------------
# Main orchestrator
# ---------------------------------------------------------------------------


def conferir_trabalho(
    registros: list[RegistroOficio],
    dados_planilha: list[list[Any]],
    q: Any,
) -> RelatorioConferencia:
    """Orquestrador da Fase 6: verifica e corrige automaticamente todos os ofícios.

    Para cada ofício gerado, executa uma conferência rigorosa em loop iterativo
    (máximo de 3 tentativas) cobrindo:
    1. Consistência de dados (ctx ↔ dados_grupo da IA).
    2. Concordância linguística (pronomes, gênero, número).
    3. Integridade do template (vazamento de tags como {{ ou }}).
    4. Linha correspondente na planilha Excel.

    As correções são aplicadas, o documento é re-renderizado e re-verificado
    até que todos os erros sejam sanados ou o limite de tentativas seja atingido.

    Args:
        registros: Lista de :class:`RegistroOficio` construída durante a Fase 3.
        dados_planilha: Lista mutável de linhas da planilha; corrigida in-place.
        q: Fila de mensagens do worker (``queue.Queue[tuple]``) para atualizar a UI.

    Returns:
        :class:`RelatorioConferencia` com o resultado completo.
    """
    from docxtpl import DocxTemplate  # noqa: PLC0415 — importado aqui para não poluir core/

    relatorio = RelatorioConferencia()
    relatorio.total_verificados = len(registros)

    q.put(("log", "\n🔍  Iniciando conferência dos ofícios gerados…", "accent"))
    logger.info(
        "Fase 6: Conferência rigorosa — %d ofício(s) a verificar.", len(registros)
    )

    for i, registro in enumerate(registros, start=1):
        resultado = ResultadoVerificacao(arquivo=registro.nome_arquivo)
        prefixo = f"  [{i}/{len(registros)}]  {registro.nome_arquivo}"

        logger.debug("Conferindo rigorosamente: %s", registro.nome_arquivo)
        q.put(("log", f"\n{prefixo}", "dim"))

        max_tentativas = 3
        tentativa = 1
        idx = registro.linha_planilha_idx

        while tentativa <= max_tentativas:
            # ── 1. Executar checagens de validação ─────────────────────────────
            resultado.erros_dados = verificar_consistencia_dados(registro)
            resultado.erros_linguisticos = verificar_concordancia_linguistica(registro)
            resultado.erros_tags = verificar_tags_pendentes(registro.caminho)

            if 0 <= idx < len(dados_planilha):
                resultado.erros_planilha = verificar_linha_planilha(
                    dados_planilha[idx], registro
                )
            else:
                resultado.erros_planilha = []

            # Se não houver nenhum erro em nenhuma categoria, validação com sucesso
            if not resultado.tem_erros:
                if tentativa > 1:
                    resultado.corrigido = True
                    logger.info("  ✔  '%s' corrigido na tentativa %d.", registro.nome_arquivo, tentativa)
                    q.put(("log", f"    ✔  Corrigido com sucesso na tentativa {tentativa}.", "success"))
                else:
                    logger.info("  ✔  '%s' OK — sem erros.", registro.nome_arquivo)
                    q.put(("log", "    ✔  OK — sem erros.", "success"))
                break

            # Se for a primeira rodada que identificou erros, incrementa o total
            if tentativa == 1:
                relatorio.total_com_erros += 1

            # Logar erros encontrados na tentativa atual
            logger.warning(
                "  [TENTATIVA %d] Erros em %s: dados=%d, ling=%d, tags=%d, planilha=%d",
                tentativa,
                registro.nome_arquivo,
                len(resultado.erros_dados),
                len(resultado.erros_linguisticos),
                len(resultado.erros_tags),
                len(resultado.erros_planilha),
            )

            q.put((
                "log",
                f"    🔧  Tentativa {tentativa}/{max_tentativas}: erros encontrados. Corrigindo e re-renderizando…",
                "warn",
            ))

            for e in resultado.erros_dados:
                q.put(("log", f"        • [DADOS] {e}", "dim"))
            for e in resultado.erros_linguisticos:
                q.put(("log", f"        • [LINGUÍSTICO] {e}", "dim"))
            for e in resultado.erros_tags:
                q.put(("log", f"        • [TEMPLATES] {e}", "dim"))
            for e in resultado.erros_planilha:
                q.put(("log", f"        • [PLANILHA] {e}", "dim"))

            # ── 2. Tentar corrigir os erros e re-renderizar ───────────────────
            try:
                # O corretor do ctx usa os erros de dados e linguísticos (incluindo tags/outros) para ajustar o ctx
                ctx_corr = corrigir_ctx(
                    registro, resultado.erros_dados, resultado.erros_linguisticos
                )

                # Se houver erros de tags pendentes na renderização do template, recarrega e limpa ctx
                if resultado.erros_tags:
                    # Garantir que re-renderiza com o ctx corrigido e limpo
                    pass

                doc = DocxTemplate(registro.template_path)
                doc.render(ctx_corr)
                doc.save(registro.caminho)

                # Atualiza o contexto no registro
                registro.ctx.update(ctx_corr)

                # Corrige a linha da planilha correspondente
                if 0 <= idx < len(dados_planilha) and resultado.erros_planilha:
                    dados_planilha[idx] = _corrigir_linha_planilha(
                        dados_planilha[idx], registro
                    )

            except Exception as exc:  # noqa: BLE001
                resultado.incorrigivel = True
                logger.error(
                    "  ✗  Falha crítica de exceção ao tentar corrigir '%s': %s",
                    registro.nome_arquivo, exc,
                )
                q.put(("log", f"    ✗  Falha crítica ao corrigir: {exc}", "error"))
                break

            tentativa += 1
        else:
            # Se o loop terminou sem "break", os erros persistiram após o limite
            resultado.incorrigivel = True
            logger.error(
                "  ✗  Não foi possível sanar todos os erros de '%s' após %d tentativas.",
                registro.nome_arquivo, max_tentativas,
            )
            q.put((
                "log",
                f"    ✗  Erro persistente: impossível sanar todos os erros após {max_tentativas} tentativas.",
                "error",
            ))

        if resultado.corrigido:
            relatorio.total_corrigidos += 1
        elif resultado.incorrigivel:
            relatorio.total_incorrigiveis += 1

        relatorio.resultados.append(resultado)

    # ── Resumo final ──────────────────────────────────────────────────────────
    _emitir_resumo(relatorio, q)
    logger.info(
        "Conferência finalizada: %d verificados, %d com erros, %d corrigidos, %d incorrigíveis.",
        relatorio.total_verificados,
        relatorio.total_com_erros,
        relatorio.total_corrigidos,
        relatorio.total_incorrigiveis,
    )
    return relatorio


def _emitir_resumo(relatorio: RelatorioConferencia, q: Any) -> None:
    """Posta na fila e no logger um resumo da rodada de conferência."""
    total = relatorio.total_verificados
    erros = relatorio.total_com_erros
    corr = relatorio.total_corrigidos
    incorr = relatorio.total_incorrigiveis

    if erros == 0:
        msg = (
            f"\n✅  Conferência concluída — {total} ofício(s) verificado(s), "
            f"nenhum erro encontrado."
        )
        q.put(("log", msg, "success"))
    else:
        partes = [
            f"{total} verificado(s)",
            f"{erros} com erro(s)",
            f"{corr} corrigido(s)",
        ]
        if incorr:
            partes.append(f"{incorr} incorrigível(is)")
        tag = "warn" if incorr else "success"
        msg = "\n📋  Conferência concluída — " + ", ".join(partes) + "."
        q.put(("log", msg, tag))
